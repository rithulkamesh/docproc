"""Word (.docx) document loader."""

from pathlib import Path
from typing import Iterator

from docproc.doc.loaders.base import DocumentLoader, LoadedPage
from docproc.doc.regions import BoundingBox, Region, RegionType


class DOCXLoader(DocumentLoader):
    """Load Word documents via python-docx."""

    def load(self, path: Path) -> Iterator[LoadedPage]:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        doc = Document(path)
        regions = []
        page_text = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                regions.append(
                    Region(
                        region_type=RegionType.TEXT,
                        bbox=BoundingBox(x1=0, y1=0, x2=0, y2=0),
                        confidence=0.95,
                        content=t,
                        metadata={"source": "paragraph"},
                    )
                )
                page_text.append(t)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if table_text.strip():
                regions.append(
                    Region(
                        region_type=RegionType.TABLE,
                        bbox=BoundingBox(x1=0, y1=0, x2=0, y2=0),
                        confidence=0.95,
                        content=table_text,
                        metadata={"source": "table"},
                    )
                )
                page_text.append(table_text)
        yield LoadedPage(
            page_num=0,
            text="\n\n".join(page_text),
            regions=regions,
            raw_images=[],
        )

    def get_full_text(self, path: Path) -> str:
        from docx import Document

        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n\n".join(parts)
